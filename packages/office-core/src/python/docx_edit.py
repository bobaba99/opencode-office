import os

from _worker import run, WorkerError
from _docx_common import iter_blocks, render_table
from _textops import para_text, replace_in_paragraph
from docx import Document

MD_STYLES = [("### ", "Heading 3"), ("## ", "Heading 2"), ("# ", "Heading 1"), ("- ", "List Bullet")]


def build_index(doc):
    return {f"{prefix}:{index}": (prefix, element) for prefix, index, element in iter_blocks(doc)}


def require(index, op_name, target, kinds):
    if target not in index:
        raise WorkerError(
            "TARGET_NOT_FOUND",
            f"No element {target}",
            "IDs come from office_read and shift as the batch applies; re-read the file, or order ops bottom-up.",
        )
    prefix, element = index[target]
    if prefix not in kinds:
        raise WorkerError(
            "BAD_TARGET_KIND",
            f"{op_name} cannot target {target} (a {prefix} element)",
            f"{op_name} targets {' or '.join(kinds)} elements.",
        )
    return prefix, element


def check_anchor(current, anchor, target):
    if anchor not in current:
        raise WorkerError(
            "ANCHOR_MISMATCH",
            f"Anchor not found at {target}",
            f"At this point in the batch the element reads: {current[:300]!r}. The batch was rolled back — nothing was written; re-read the file and retry with corrected ops.",
        )


def styled_line(line):
    for marker, style in MD_STYLES:
        if line.startswith(marker):
            return line[len(marker):], style
    return line, None


def add_styled_paragraph(doc, text, style):
    try:
        return doc.add_paragraph(text, style=style) if style else doc.add_paragraph(text)
    except KeyError:
        raise WorkerError(
            "STYLE_NOT_FOUND",
            f"Style {style!r} does not exist in this document",
            "Only styles the document defines can be used; # / ## / ### / - map to Heading 1-3 / List Bullet.",
        )


def apply_one(doc, op):
    kind = op["op"]
    index = build_index(doc)
    if kind == "replace_text":
        _, el = require(index, kind, op["target"], ["p"])
        current = para_text(el)
        check_anchor(current, op["anchor"], op["target"])
        if current.count(op["anchor"]) > 1:
            raise WorkerError(
                "AMBIGUOUS_ANCHOR",
                f"Anchor occurs more than once in {op['target']}",
                "Extend the anchor with surrounding text until it is unique within the element.",
            )
        replace_in_paragraph(el, op["anchor"], op["text"])
        return {"op": kind, "target": op["target"], "text_after": para_text(el)}
    if kind == "insert_content":
        _, el = require(index, kind, op["after"], ["p", "tbl"])
        anchor_el = el._p if hasattr(el, "_p") else el._tbl
        for line in [line for line in op["markdown"].splitlines() if line.strip()]:
            text, style = styled_line(line)
            new_p = add_styled_paragraph(doc, text, style)
            anchor_el.addnext(new_p._p)
            anchor_el = new_p._p
        return {"op": kind, "after": op["after"]}
    if kind == "delete_element":
        prefix, el = require(index, kind, op["target"], ["p", "tbl"])
        current = para_text(el) if prefix == "p" else render_table(el).split("\n")[0]
        check_anchor(current, op["anchor"], op["target"])
        xml_el = el._p if prefix == "p" else el._tbl
        xml_el.getparent().remove(xml_el)
        return {"op": kind, "target": op["target"]}
    if kind == "set_style":
        _, el = require(index, kind, op["target"], ["p"])
        current = para_text(el)
        check_anchor(current, op["anchor"], op["target"])
        try:
            el.style = doc.styles[op["style"]]
        except KeyError:
            raise WorkerError(
                "STYLE_NOT_FOUND",
                f"Style {op['style']!r} does not exist in this document",
                "office_read shows each paragraph's style; only styles the document defines can be applied.",
            )
        return {"op": kind, "target": op["target"]}
    if kind == "set_table_cell":
        _, el = require(index, kind, op["target"], ["tbl"])
        rows, cols = len(el.rows), len(el.columns)
        if not (0 <= op["row"] < rows and 0 <= op["col"] < cols):
            raise WorkerError(
                "CELL_OUT_OF_RANGE",
                f"{op['target']} is {rows}x{cols}; cell ({op['row']},{op['col']}) does not exist",
                "Row and col are 0-based and must be inside the dimensions office_read reports.",
            )
        if op.get("anchor") is not None and op["anchor"] != el.cell(op["row"], op["col"]).text:
            raise WorkerError(
                "ANCHOR_MISMATCH",
                f"Cell ({op['row']},{op['col']}) of {op['target']} does not match anchor",
                f"Cell currently reads: {el.cell(op['row'], op['col']).text[:300]!r}. The batch was rolled back — nothing was written; re-read the file and retry with corrected ops.",
            )
        el.cell(op["row"], op["col"]).text = op["text"]
        return {"op": kind, "target": op["target"]}
    raise WorkerError(
        "UNKNOWN_OP",
        f"Unknown docx op: {kind}",
        "Valid ops: replace_text, insert_content, delete_element, set_style, set_table_cell.",
    )


def main(payload):
    path = payload["file"]
    try:
        doc = Document(path)
    except Exception as e:
        raise WorkerError("FILE_OPEN", f"Could not open {path} as .docx: {e}", "Check the path; the file must be a .docx (not legacy .doc).")

    results = [apply_one(doc, op) for op in payload["operations"]]
    tmp = path + ".tmp-opencode-office"
    try:
        doc.save(tmp)
        os.replace(tmp, path)
    finally:
        if os.path.exists(tmp):
            os.remove(tmp)
    return {"applied": len(results), "results": results}


run(main)
